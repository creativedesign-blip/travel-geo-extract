"""Read .txt/.md/.docx/.pdf and emit plain text to stdout.

Usage:
    python read_doc.py path/to/file.pdf

Dependencies (installed lazily — only the library for the input format is needed):
    .docx       -> python-docx (`pip install python-docx`)
    .pdf        -> pdfplumber (`pip install pdfplumber`), fallback pypdf
    .txt / .md  -> stdlib only
"""

from __future__ import annotations

import sys
from pathlib import Path


def read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="replace")


def read_docx(path: Path) -> str:
    try:
        from docx import Document  # python-docx
    except ImportError as e:
        raise SystemExit(
            "python-docx not installed. Run: pip install python-docx"
        ) from e
    doc = Document(str(path))
    parts = []
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def read_pdf(path: Path) -> str:
    # try pdfplumber first (better text extraction)
    try:
        import pdfplumber  # type: ignore
        parts = []
        with pdfplumber.open(str(path)) as pdf:
            for page in pdf.pages:
                txt = page.extract_text() or ""
                if txt.strip():
                    parts.append(txt)
        return "\n".join(parts)
    except ImportError:
        pass
    # fallback to pypdf
    try:
        from pypdf import PdfReader  # type: ignore
    except ImportError:
        try:
            from PyPDF2 import PdfReader  # type: ignore
        except ImportError as e:
            raise SystemExit(
                "No PDF library found. Run: pip install pdfplumber"
            ) from e
    reader = PdfReader(str(path))
    parts = [page.extract_text() or "" for page in reader.pages]
    return "\n".join(p for p in parts if p.strip())


def read_any(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix in {".txt", ".md", ".markdown"}:
        return read_text(path)
    if suffix == ".docx":
        return read_docx(path)
    if suffix == ".pdf":
        return read_pdf(path)
    raise SystemExit(f"Unsupported file type: {suffix}")


def main() -> None:
    if len(sys.argv) != 2:
        print("usage: read_doc.py <path>", file=sys.stderr)
        sys.exit(2)
    path = Path(sys.argv[1])
    if not path.exists():
        print(f"file not found: {path}", file=sys.stderr)
        sys.exit(2)
    sys.stdout.write(read_any(path))


if __name__ == "__main__":
    main()
